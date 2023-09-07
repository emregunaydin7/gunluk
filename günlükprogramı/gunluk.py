import tkinter as tk
from tkcalendar import Calendar
from docx import Document
import os
from datetime import datetime

# Kullanıcı bilgilerini tutmak için bir sözlük oluşturun
users = {"username": "password"}

# Kaydedilen belgelerin tarihlerini tutmak için bir liste oluşturun
saved_dates = []

def load_saved_dates():
    # Kaydedilmiş belgelerin tarihlerini mevcut klasörde arayarak Listbox'a ekleyin
    for file_name in os.listdir():
        if file_name.startswith("TarihRaporu_") and file_name.endswith(".docx"):
            date_str = file_name.replace("TarihRaporu_", "").replace(".docx", "")
            if date_str not in saved_dates:
                saved_dates.append(date_str)
                selected_dates_listbox.insert(tk.END, date_str)

    # Tarihleri sıralayın
    saved_dates.sort(key=lambda x: datetime.strptime(x, "%d-%m-%Y"))

def get_selected_date():
    selected_date = cal.get_date()
    selected_date = datetime.strptime(selected_date, "%d/%m/%Y")  # Tarihi bir datetime nesnesine çevir
    selected_date_str = selected_date.strftime("%d-%m-%Y")  # Tarihi geçerli bir formata dönüştür
    selected_date_label.config(text=f"Seçilen Tarih: {selected_date_str}")
    open_or_create_word_document(selected_date_str)
    add_to_listbox(selected_date_str)  # Seçilen tarihi Listbox'a ekleyin

def open_or_create_word_document(selected_date):
    document_name = f"TarihRaporu_{selected_date}.docx"

    if os.path.exists(document_name):
        open_word_document(document_name)
    else:
        create_word_document(selected_date)
        open_word_document(document_name)

def create_word_document(selected_date):
    doc = Document()
    doc.add_heading("Tarih Raporu", 0)
    doc.add_paragraph(f"Seçilen Tarih: {selected_date}")
    document_name = f"TarihRaporu_{selected_date}.docx"
    doc.save(document_name)

def open_word_document(document_name):
    os.system(f"start {document_name}")

def add_to_listbox(selected_date):
    if selected_date not in saved_dates:
        saved_dates.append(selected_date)
        selected_dates_listbox.insert(tk.END, selected_date)  # Seçilen tarihi Listbox'a ekleyin

def login():
    username = username_entry.get()
    password = password_entry.get()

    # Kullanıcı adı ve şifre doğrulamasını kontrol edin
    if username in users and users[username] == password:
        result_label.config(text="Giriş başarılı!")
       
        root.geometry("600x600")
        username_label.pack_forget()
        username_entry.pack_forget()
        password_label.pack_forget()
        password_entry.pack_forget()
        login_button.pack_forget()
    else:
        result_label.config(text="Giriş başarısız. Lütfen kullanıcı adı ve şifrenizi kontrol edin.")



root = tk.Tk()
root.title("GÜNLÜK")
root.geometry("250x150")
root.iconbitmap("icon.ico")


# Kullanıcı adı etiketi ve giriş alanı
username_label = tk.Label(root, text="Kullanıcı Adı:")
username_label.pack()
username_entry = tk.Entry(root)
username_entry.pack()

# Şifre etiketi ve giriş alanı
password_label = tk.Label(root, text="Şifre:")
password_label.pack()
password_entry = tk.Entry(root, show="*")
password_entry.pack()

# Giriş düğmesi
login_button = tk.Button(root, text="Giriş Yap", command=login)
login_button.pack()

# Sonuç etiketi
result_label = tk.Label(root, text="")
result_label.pack()

# Takvim bileşeni
cal = Calendar(root, selectmode="day", date_pattern="dd/MM/yyyy")
cal.pack(pady=20)

# Tarih seçme düğmesi
select_date_button = tk.Button(root, text="Tarihi Seç", command=get_selected_date)
select_date_button.pack()

# Seçilen tarih etiketi
selected_date_label = tk.Label(root, text="")
selected_date_label.pack()

# Listbox'ta seçilen tarihleri görüntülemek için
selected_dates_listbox = tk.Listbox(root)
selected_dates_listbox.pack()

# Kaydedilmiş tarihleri yükleyin
load_saved_dates()

root.mainloop()
